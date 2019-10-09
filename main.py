import os
import re
from os import path

import mglearn
import numpy as np
import pandas as pd
# %pylab
# %matplotlib inline
from matplotlib import pyplot as plt
from matplotlib import pyplot as plt_2
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from wordcloud import WordCloud, STOPWORDS
from gensim.summarization import keywords

import global_var

stopwords = set(STOPWORDS)
import pdf_2_txt
import summarizer
from docx import Document
from gensim.summarization import summarize

def main(filePath,filename):
    vectorizer = TfidfVectorizer()
    sourcefile=filePath+'/'+filename+'.pdf'
    lone=pdf_2_txt.convert_pdf_to_txt(sourcefile)
    clean_cont = lone.splitlines()

    shear=[i.replace('\xe2\x80\x9c','') for i in clean_cont ]
    shear=[i.replace('\xe2\x80\x9d','') for i in shear ]
    shear=[i.replace('\xe2\x80\x99s','') for i in shear ]

    shears = [x for x in shear if x != ' ']
    shearss = [x for x in shears if x != '']
    dubby=[re.sub("[^a-zA-Z]+", " ", s) for s in shearss]


    vect=CountVectorizer(ngram_range=(1,1),stop_words='english')
    dtm=vect.fit_transform(dubby)
    pd.DataFrame(dtm.toarray(),columns=vect.get_feature_names())

    #lda
    lda=LatentDirichletAllocation(n_components=5)
    lda_dtf=lda.fit_transform(dtm)

    sorting=np.argsort(lda.components_)[:,::-1]
    features=np.array(vect.get_feature_names())

    mglearn.tools.print_topics(topics=range(5), feature_names=features,
    sorting=sorting, topics_per_chunk=5, n_words=10)

    wordcloud = WordCloud(width = 800, height = 800,
                    background_color ='white',
                    stopwords = stopwords,
                    min_font_size = 10).generate(lone)

    # plot the WordCloud image
    plt.figure()
    plt.imshow(wordcloud)
    plt.axis("off")
    plt.tight_layout(pad=0)

    #plt.show()
    plt.savefig('word_cloud_graph/N.png')
    plt.clf()
    X = vectorizer.fit_transform(dubby)
    key_words=keywords(lone,words = 10,scores = True, lemmatize = True)
    key_data=str(key_words)
    summary=summarize(lone)

    # chech if dir path exists
    filePath = filePath + '/' + 'reports'
    if not (os.path.exists(filePath)):
        os.mkdir(filePath)

    # set report file name
    filePath = filePath + '/' + filename + ".docx"


    document = Document()

    document.add_heading('Summarizing the document', level=1)

    strng=summarizer.generate_Summary(lone)
    clean_cont = strng.splitlines()

    shear=[i.replace('\xe2\x80\x9c','') for i in clean_cont ]
    shear=[i.replace('\xe2\x80\x9d','') for i in shear ]
    shear=[i.replace('\xe2\x80\x99s','') for i in shear ]

    shears = [x for x in shear if x != ' ']
    shearss = [x for x in shears if x != '']
    dubby=[re.sub("[^a-zA-Z]+", " ", s) for s in shearss]
    #document.add_paragraph(dubby)
    print(key_words)
    print(type(key_words))

    #key_data=''.join(keywords)
    print(type(keywords))
    print(type(key_data))

    plt_2.figure()
    plt_2.bar(range(len(key_words)), [val[1] for val in key_words])
    plt_2.xticks(range(len(key_words)), [val[0] for val in key_words])
    plt_2.xticks(rotation=70)
    plt.tight_layout(pad=0)
    plt_2.autoscale()
    #plt_2.show()
    plt_2.savefig('plot.png')
    document.add_heading('Topics in... ', level=2)
    document.add_paragraph(key_data)
    document.add_picture('plot.png')

    document.add_heading('Summary...', level=2)
    document.add_picture('word_cloud_graph/N.png')
    document.add_paragraph(summary)

    document.save(filePath)
    global_var.reportfilepath=filePath
    # Domain_Name_Topic=np.argsort(lda_dtf[:,4])[::-1]
    # for i in Domain_Name_Topic[:4]:
    #     print(".".join(dubby[i].split(".")[:2]) + ".\n")
    #
    # Agreement_Topic=np.argsort(lda_dtf[:,2])[::-1]
    # for i in Agreement_Topic[:4]:
    #     print(".".join(dubby[i].split(".")[:2]) + ".\n")
